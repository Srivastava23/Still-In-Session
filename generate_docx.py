"""
Generate Word Document for Still-In-Session Project
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('STILL-IN-SESSION')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB) # Blue color

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('HTTP Session Management Demo')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

tech = doc.add_paragraph()
tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tech.add_run('Java Jakarta Servlets Application')
run.font.size = Pt(13)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Submitted by')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = name_p.add_run('Ananya Chandra')
run.bold = True
run.font.size = Pt(16)

reg = doc.add_paragraph()
reg.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = reg.add_run('Registration No: 25MCA1065')
run.font.size = Pt(13)

doc.add_page_break()

# ============================================================
# AIM
# ============================================================
doc.add_heading('1. AIM', level=1)
doc.add_paragraph(
    'To develop a minimal web application "Still-In-Session" using Jakarta Servlets that '
    'demonstrates the complete lifecycle of HTTP Session Management, including session creation, '
    'attribute storage, attribute retrieval, and session invalidation across multiple requests.'
)

# ============================================================
# TOOLS
# ============================================================
doc.add_heading('2. TOOLS & TECHNOLOGIES', level=1)

table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
data = [
    ['Language', 'Java 17'],
    ['Web API', 'Jakarta Servlet 6.0'],
    ['Server', 'Embedded Tomcat 10.1.19'],
    ['Frontend', 'HTML5'],
    ['Build Tool', 'Gradle (Gretty Plugin)'],
]
for i, h in enumerate(['Component', 'Technology']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph()

# ============================================================
# PROCEDURE (Brief)
# ============================================================
doc.add_heading('3. PROCEDURE', level=1)

steps = [
    'Set up Gradle project with the Gretty plugin for embedded Tomcat 10 and Jakarta Servlet API 6.0.',
    'Created a simple HTML login form (login.html) to capture a username and POST it to the server.',
    'Developed LoginServlet.java handles the POST request, retrieves the username parameter, creates a new HTTP session (request.getSession()), stores the username as a session attribute, and redirects.',
    'Developed WelcomeServlet.java retrieves the existing session (request.getSession(false)), reads the username attribute, and renders a dynamic welcome page showing the Session ID.',
    'Developed LogoutServlet.java retrieves the existing session, calls session.invalidate() to destroy all session data, and displays a logout confirmation.',
    'Configured web.xml deployment descriptor to map URL patterns (/login, /welcome, /logout) to their respective servlet classes.',
    'Tested the complete flow (Login -> Welcome -> Logout) locally on embedded Tomcat.',
]

for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()

# ============================================================
# KEY SESSION METHODS
# ============================================================
doc.add_heading('4. KEY SESSION METHODS USED', level=1)

table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Method', 'Purpose']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
db_rows = [
    ['request.getSession()', 'Creates a new session or retrieves an existing one.'],
    ['request.getSession(false)', 'Retrieves existing session only (returns null if none exists).'],
    ['session.setAttribute("key", value)', 'Stores an object within the session state.'],
    ['session.getAttribute("key")', 'Retrieves a previously stored object from the session.'],
    ['session.invalidate()', 'Destroys the session and unbinds all its objects.'],
    ['session.getId()', 'Returns a string containing the unique identifier assigned to this session.'],
]
for i, row in enumerate(db_rows):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_page_break()

# ============================================================
# KEY SOURCE CODE
# ============================================================
doc.add_heading('5. SOURCE CODE', level=1)

PROJECT_DIR = r'e:\D Drive\Documents\Projects\Still-In-Session'

code_files = [
    ('5.1 LoginServlet.java (Session Creation)', 'src/main/java/com/session/LoginServlet.java'),
    ('5.2 WelcomeServlet.java (Session Usage)', 'src/main/java/com/session/WelcomeServlet.java'),
    ('5.3 LogoutServlet.java (Session Invalidation)', 'src/main/java/com/session/LogoutServlet.java'),
]

for heading, filepath in code_files:
    doc.add_heading(heading, level=2)
    full_path = os.path.join(PROJECT_DIR, filepath)
    try:
        with open(full_path, 'r', encoding='utf-8') as f:
            code = f.read()
        p = doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(7.5)
    except Exception as e:
        doc.add_paragraph(f'[Error: {e}]')
    doc.add_page_break()

# ============================================================
# OUTPUT SCREENSHOTS
# ============================================================
doc.add_heading('6. OUTPUT SCREENSHOTS', level=1)

screenshots_dir = r'C:\Users\Anany\.gemini\antigravity\brain\76453d18-1496-4593-9db8-3edd4b77856a'

screenshot_info = [
    ('6.1 Login Page (/SessionDemo/)', 'login_page', 'Figure 1: HTML Login Form taking username input.'),
    ('6.2 Welcome Page (/welcome)', 'welcome_page', 'Figure 2: Dynamic Welcome page showing retrieved username and active Session ID.'),
    ('6.3 Logout Page (/logout)', 'logout_page', 'Figure 3: Confirmation page after session.invalidate() was called.')
]

for heading, img_prefix, caption in screenshot_info:
    doc.add_heading(heading, level=2)
    img_added = False
    for f in os.listdir(screenshots_dir):
        if img_prefix in f and f.endswith('.png'):
            doc.add_picture(os.path.join(screenshots_dir, f), width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(10)
            img_added = True
            break
    if not img_added:
        doc.add_paragraph(f"[Screenshot '{img_prefix}' not found]")
    if heading != screenshot_info[-1][0]:
        doc.add_paragraph()

doc.add_page_break()

# ============================================================
# CONCLUSION
# ============================================================
doc.add_heading('7. CONCLUSION', level=1)
doc.add_paragraph(
    'The "Still-In-Session" application was successfully implemented and tested on an embedded Tomcat server. '
    'It clearly demonstrated the ability of HTTP Sessions to overcome the stateless nature of the HTTP protocol '
    'by persisting user-specific data (username) across consecutive page requests, and properly clearing that '
    'data upon logout. The project met all functional requirements for a servlet-based session management lab.'
)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(PROJECT_DIR, 'Still_In_Session_Report_Ananya_Chandra_25MCA1065.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
